"""
Post-Processing Module for PDF to DOCX Conversion

This module provides:
- Content validation and comparison
- Structure cleanup and correction
- Blank page removal
- Formatting normalization
"""

import os
import re
from typing import List, Tuple, Optional, Dict, Any
from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


@dataclass
class ContentComparison:
    """Detailed comparison between PDF and DOCX content."""
    pdf_text: str
    docx_text: str
    pdf_word_count: int
    docx_word_count: int
    missing_words: List[str]
    extra_words: List[str]
    match_percentage: float
    page_comparisons: List[Dict[str, Any]]


@dataclass
class CleanupResult:
    """Result of post-processing cleanup."""
    success: bool
    original_paragraphs: int
    final_paragraphs: int
    blank_paragraphs_removed: int
    duplicate_spaces_fixed: int
    issues_fixed: List[str]


class PostProcessor:
    """
    Post-processor for DOCX documents.

    Provides cleanup and validation functionality:
    - Remove blank paragraphs and pages
    - Fix formatting issues
    - Validate against original PDF
    - Merge split paragraphs
    """

    def __init__(self, verbose: bool = True):
        self.verbose = verbose

    def _log(self, message: str):
        if self.verbose:
            print(message)

    def compare_content(
        self,
        pdf_path: str,
        docx_path: str
    ) -> ContentComparison:
        """
        Compare content between original PDF and converted DOCX.

        Args:
            pdf_path: Path to original PDF
            docx_path: Path to converted DOCX

        Returns:
            ContentComparison with detailed analysis
        """
        self._log("Comparing PDF and DOCX content...")

        # Extract PDF text page by page
        pdf_doc = fitz.open(pdf_path)
        pdf_pages_text = []
        pdf_full_text = ""

        for page in pdf_doc:
            page_text = page.get_text()
            pdf_pages_text.append(page_text)
            pdf_full_text += page_text + "\n"

        pdf_doc.close()

        # Extract DOCX text
        docx_doc = Document(docx_path)
        docx_text = ""

        for para in docx_doc.paragraphs:
            docx_text += para.text + "\n"

        for table in docx_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    docx_text += cell.text + "\n"

        # Normalize and compare
        pdf_words = self._extract_words(pdf_full_text)
        docx_words = self._extract_words(docx_text)

        pdf_word_set = set(pdf_words)
        docx_word_set = set(docx_words)

        missing = list(pdf_word_set - docx_word_set)
        extra = list(docx_word_set - pdf_word_set)

        if pdf_word_set:
            match_pct = len(pdf_word_set & docx_word_set) / len(pdf_word_set) * 100
        else:
            match_pct = 100.0 if not docx_word_set else 0.0

        # Per-page comparison
        page_comparisons = []
        for i, page_text in enumerate(pdf_pages_text):
            page_words = set(self._extract_words(page_text))
            page_in_docx = len(page_words & docx_word_set) / len(page_words) * 100 if page_words else 100
            page_comparisons.append({
                "page": i + 1,
                "pdf_words": len(page_words),
                "match_percentage": page_in_docx
            })

        return ContentComparison(
            pdf_text=pdf_full_text,
            docx_text=docx_text,
            pdf_word_count=len(pdf_words),
            docx_word_count=len(docx_words),
            missing_words=missing[:50],  # Limit for display
            extra_words=extra[:50],
            match_percentage=match_pct,
            page_comparisons=page_comparisons
        )

    def cleanup_document(
        self,
        docx_path: str,
        output_path: Optional[str] = None,
        remove_blank_paragraphs: bool = True,
        fix_spacing: bool = True,
        normalize_fonts: bool = True
    ) -> CleanupResult:
        """
        Clean up a DOCX document.

        Args:
            docx_path: Path to DOCX to clean
            output_path: Output path (default: overwrite input)
            remove_blank_paragraphs: Remove empty paragraphs
            fix_spacing: Fix duplicate spaces
            normalize_fonts: Normalize font inconsistencies

        Returns:
            CleanupResult with cleanup statistics
        """
        self._log("Cleaning up document...")

        doc = Document(docx_path)
        original_count = len(doc.paragraphs)
        blank_removed = 0
        spaces_fixed = 0
        issues_fixed = []

        # Process paragraphs
        paragraphs_to_remove = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text

            # Track blank paragraphs
            if remove_blank_paragraphs and not text.strip():
                # Check if this is a structural blank (page break) or just empty
                if not self._is_structural_element(para):
                    paragraphs_to_remove.append(para)
                    blank_removed += 1
                continue

            # Fix spacing
            if fix_spacing and '  ' in text:
                new_text = re.sub(r'  +', ' ', text)
                if new_text != text:
                    # Update paragraph text
                    for run in para.runs:
                        run.text = re.sub(r'  +', ' ', run.text)
                    spaces_fixed += 1

        # Remove blank paragraphs (must do in reverse to maintain indices)
        for para in reversed(paragraphs_to_remove):
            self._remove_paragraph(para)

        # Clean up tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if fix_spacing:
                            for run in para.runs:
                                if '  ' in run.text:
                                    run.text = re.sub(r'  +', ' ', run.text)
                                    spaces_fixed += 1

        if blank_removed > 0:
            issues_fixed.append(f"Removed {blank_removed} blank paragraphs")
        if spaces_fixed > 0:
            issues_fixed.append(f"Fixed {spaces_fixed} spacing issues")

        # Save
        if output_path is None:
            output_path = docx_path

        doc.save(output_path)

        return CleanupResult(
            success=True,
            original_paragraphs=original_count,
            final_paragraphs=len(Document(output_path).paragraphs),
            blank_paragraphs_removed=blank_removed,
            duplicate_spaces_fixed=spaces_fixed,
            issues_fixed=issues_fixed
        )

    def remove_blank_pages(self, docx_path: str, output_path: Optional[str] = None) -> int:
        """
        Remove blank pages from document.

        Args:
            docx_path: Path to DOCX
            output_path: Output path

        Returns:
            Number of blank elements removed
        """
        doc = Document(docx_path)
        removed = 0

        # Track consecutive empty elements
        consecutive_empty = 0
        elements_to_remove = []

        for para in doc.paragraphs:
            if not para.text.strip():
                consecutive_empty += 1
                if consecutive_empty > 2:  # More than 2 consecutive blanks likely a blank page
                    elements_to_remove.append(para)
            else:
                consecutive_empty = 0

        for para in reversed(elements_to_remove):
            self._remove_paragraph(para)
            removed += 1

        if output_path is None:
            output_path = docx_path

        doc.save(output_path)
        return removed

    def validate_and_fix(
        self,
        pdf_path: str,
        docx_path: str,
        output_path: Optional[str] = None,
        min_match_threshold: float = 0.85
    ) -> Tuple[bool, ContentComparison, Optional[CleanupResult]]:
        """
        Validate conversion and attempt to fix issues.

        Args:
            pdf_path: Original PDF path
            docx_path: Converted DOCX path
            output_path: Output path for fixed document
            min_match_threshold: Minimum acceptable match percentage

        Returns:
            Tuple of (is_valid, comparison, cleanup_result)
        """
        # Initial comparison
        comparison = self.compare_content(pdf_path, docx_path)

        self._log(f"Content match: {comparison.match_percentage:.1f}%")
        self._log(f"PDF words: {comparison.pdf_word_count}, DOCX words: {comparison.docx_word_count}")

        cleanup_result = None

        # If match is below threshold, try cleanup
        if comparison.match_percentage < min_match_threshold * 100:
            self._log("Match below threshold, attempting cleanup...")

            cleanup_result = self.cleanup_document(
                docx_path,
                output_path or docx_path,
                remove_blank_paragraphs=True,
                fix_spacing=True
            )

            # Re-compare after cleanup
            comparison = self.compare_content(pdf_path, output_path or docx_path)
            self._log(f"After cleanup: {comparison.match_percentage:.1f}%")

        is_valid = comparison.match_percentage >= min_match_threshold * 100

        return is_valid, comparison, cleanup_result

    def print_comparison_report(self, comparison: ContentComparison):
        """Print a detailed comparison report."""
        print("\n" + "="*60)
        print("CONTENT COMPARISON REPORT")
        print("="*60)

        print(f"\nOverall Match: {comparison.match_percentage:.1f}%")
        print(f"PDF Word Count: {comparison.pdf_word_count}")
        print(f"DOCX Word Count: {comparison.docx_word_count}")
        print(f"Word Difference: {comparison.pdf_word_count - comparison.docx_word_count}")

        print("\nPer-Page Analysis:")
        for page_data in comparison.page_comparisons:
            status = "OK" if page_data["match_percentage"] >= 85 else "CHECK"
            print(f"  Page {page_data['page']}: {page_data['match_percentage']:.1f}% ({page_data['pdf_words']} words) [{status}]")

        if comparison.missing_words:
            print(f"\nSample Missing Words ({len(comparison.missing_words)} shown):")
            print(f"  {', '.join(comparison.missing_words[:20])}")

        print("="*60)

    def _extract_words(self, text: str) -> List[str]:
        """Extract and normalize words from text."""
        # Remove special characters and normalize
        text = re.sub(r'[^\w\s]', ' ', text)
        text = text.lower()
        words = text.split()
        # Filter very short words and numbers
        words = [w for w in words if len(w) > 2 and not w.isdigit()]
        return words

    def _is_structural_element(self, para) -> bool:
        """Check if paragraph is a structural element (like page break)."""
        # Check for page breaks in runs
        for run in para.runs:
            if run._element.xml.find('w:br') != -1:
                return True
        return False

    def _remove_paragraph(self, para):
        """Remove a paragraph from document."""
        p = para._element
        p.getparent().remove(p)


class ConversionValidator:
    """
    Validates PDF to DOCX conversion quality.
    """

    def __init__(self):
        self.post_processor = PostProcessor(verbose=False)

    def full_validation(
        self,
        pdf_path: str,
        docx_path: str
    ) -> Dict[str, Any]:
        """
        Perform full validation of conversion.

        Returns detailed validation report.
        """
        report = {
            "pdf_path": pdf_path,
            "docx_path": docx_path,
            "valid": False,
            "content_match": 0.0,
            "structure_check": {},
            "issues": [],
            "recommendations": []
        }

        # Content comparison
        comparison = self.post_processor.compare_content(pdf_path, docx_path)
        report["content_match"] = comparison.match_percentage

        # Structure checks
        pdf_doc = fitz.open(pdf_path)
        docx_doc = Document(docx_path)

        # Page count comparison (estimate DOCX pages)
        pdf_pages = len(pdf_doc)
        docx_paragraphs = len(docx_doc.paragraphs)

        report["structure_check"] = {
            "pdf_pages": pdf_pages,
            "docx_paragraphs": docx_paragraphs,
            "pdf_word_count": comparison.pdf_word_count,
            "docx_word_count": comparison.docx_word_count
        }

        # Check for issues
        if comparison.match_percentage < 85:
            report["issues"].append("Content match below 85% threshold")
            report["recommendations"].append("Review complex layouts or tables")

        if comparison.docx_word_count < comparison.pdf_word_count * 0.8:
            report["issues"].append("Significant content loss detected")
            report["recommendations"].append("Check for scanned/image-based pages")

        if docx_paragraphs < pdf_pages:
            report["issues"].append("Document may be missing pages")
            report["recommendations"].append("Check for conversion errors")

        # Check for blank page issues
        blank_paras = sum(1 for p in docx_doc.paragraphs if not p.text.strip())
        if blank_paras > pdf_pages * 2:
            report["issues"].append(f"Excessive blank paragraphs ({blank_paras})")
            report["recommendations"].append("Run cleanup to remove blank content")

        pdf_doc.close()

        # Overall validity
        report["valid"] = (
            comparison.match_percentage >= 85 and
            len(report["issues"]) == 0
        )

        return report

    def print_validation_report(self, report: Dict[str, Any]):
        """Print formatted validation report."""
        print("\n" + "="*60)
        print("CONVERSION VALIDATION REPORT")
        print("="*60)

        status = "PASSED" if report["valid"] else "FAILED"
        print(f"\nStatus: {status}")
        print(f"Content Match: {report['content_match']:.1f}%")

        print(f"\nStructure:")
        for key, value in report["structure_check"].items():
            print(f"  {key}: {value}")

        if report["issues"]:
            print(f"\nIssues Found ({len(report['issues'])}):")
            for issue in report["issues"]:
                print(f"  - {issue}")

        if report["recommendations"]:
            print(f"\nRecommendations:")
            for rec in report["recommendations"]:
                print(f"  - {rec}")

        print("="*60)

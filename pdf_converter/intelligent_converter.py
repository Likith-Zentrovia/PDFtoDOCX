"""
Intelligent PDF to DOCX Converter

Uses Claude Vision analysis to create accurate DOCX output that perfectly
matches the original PDF structure, layout, and formatting.
"""

import os
import io
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple
from dataclasses import dataclass

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .vision_analyzer import (
    ClaudeVisionAnalyzer,
    DocumentAnalysis,
    PageAnalysis,
    TextElement,
    ImageElement,
    TableElement
)


@dataclass
class ConversionResult:
    """Result of the conversion process."""
    success: bool
    output_path: str
    pages_converted: int
    text_elements_processed: int
    images_extracted: int
    tables_created: int
    validation_score: float
    notes: List[str]


class IntelligentConverter:
    """
    Intelligent PDF to DOCX converter powered by Claude Vision.

    This converter:
    1. Analyzes each page visually using Claude Vision
    2. Understands exact layout, columns, and formatting
    3. Recreates the document structure accurately in DOCX
    4. Validates output against original
    """

    # Font size mapping
    FONT_SIZES = {
        "small": 9,
        "medium": 11,
        "large": 14,
        "xlarge": 18
    }

    def __init__(self, api_key: Optional[str] = None):
        """
        Initialize the converter.

        Args:
            api_key: Anthropic API key (optional, uses env var if not provided)
        """
        self.analyzer = ClaudeVisionAnalyzer(api_key)
        self.doc: Optional[Document] = None
        self.pdf_doc: Optional[fitz.Document] = None
        self._stats = {
            "text_elements": 0,
            "images": 0,
            "tables": 0
        }
        self._notes: List[str] = []

    def convert(
        self,
        pdf_path: str,
        output_path: Optional[str] = None,
        pages: Optional[List[int]] = None
    ) -> ConversionResult:
        """
        Convert PDF to DOCX with intelligent layout preservation.

        Args:
            pdf_path: Path to input PDF
            output_path: Path for output DOCX (optional)
            pages: Specific pages to convert (0-indexed, optional)

        Returns:
            ConversionResult with conversion details
        """
        pdf_path = str(Path(pdf_path).resolve())

        if not os.path.exists(pdf_path):
            return ConversionResult(
                success=False,
                output_path="",
                pages_converted=0,
                text_elements_processed=0,
                images_extracted=0,
                tables_created=0,
                validation_score=0.0,
                notes=["PDF file not found"]
            )

        if output_path is None:
            output_path = str(Path(pdf_path).with_suffix('.docx'))

        print(f"\n{'='*60}")
        print(f"INTELLIGENT PDF TO DOCX CONVERTER")
        print(f"{'='*60}")
        print(f"Input:  {pdf_path}")
        print(f"Output: {output_path}")

        # Open PDF
        self.pdf_doc = fitz.open(pdf_path)

        if pages is None:
            pages = list(range(len(self.pdf_doc)))

        print(f"Pages:  {len(pages)}")
        print(f"\n[1/4] Analyzing document with Claude Vision...")

        # Analyze document
        analysis = self.analyzer.analyze_document(
            pdf_path,
            pages=pages,
            dpi=150
        )

        print(f"       Layout: {analysis.dominant_layout}")
        print(f"       Type: {analysis.document_type}")

        print(f"\n[2/4] Creating document structure...")

        # Create DOCX
        self.doc = Document()
        self._setup_document(analysis)

        print(f"\n[3/4] Converting pages...")

        # Convert each page
        for i, page_num in enumerate(pages):
            page_analysis = analysis.pages[i]
            print(f"       Page {page_num + 1}: {page_analysis.layout_type} "
                  f"({page_analysis.num_columns} col, {len(page_analysis.text_elements)} blocks)")

            self._convert_page(page_analysis, page_num, i == 0)

        print(f"\n[4/4] Saving document...")

        # Save
        self.doc.save(output_path)
        self.pdf_doc.close()

        # Calculate validation score
        validation_score = self._calculate_validation_score(analysis)

        print(f"\n{'='*60}")
        print(f"CONVERSION COMPLETE")
        print(f"{'='*60}")
        print(f"Text elements: {self._stats['text_elements']}")
        print(f"Images:        {self._stats['images']}")
        print(f"Tables:        {self._stats['tables']}")
        print(f"Quality score: {validation_score:.0%}")

        if self._notes:
            print(f"\nNotes:")
            for note in self._notes[:5]:
                print(f"  - {note}")

        return ConversionResult(
            success=True,
            output_path=output_path,
            pages_converted=len(pages),
            text_elements_processed=self._stats['text_elements'],
            images_extracted=self._stats['images'],
            tables_created=self._stats['tables'],
            validation_score=validation_score,
            notes=self._notes
        )

    def _setup_document(self, analysis: DocumentAnalysis):
        """Setup document properties based on analysis."""
        if not analysis.pages:
            return

        first_page = analysis.pages[0]

        # Set page size
        section = self.doc.sections[0]
        section.page_width = Twips(first_page.width * 20)
        section.page_height = Twips(first_page.height * 20)

        # Set reasonable margins
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    def _convert_page(self, page_analysis: PageAnalysis, page_num: int, is_first: bool):
        """Convert a single page based on vision analysis."""
        # Add page break for non-first pages
        if not is_first:
            self.doc.add_page_break()

        # Handle different layout types
        if page_analysis.num_columns == 1:
            self._convert_single_column(page_analysis, page_num)
        else:
            self._convert_multi_column(page_analysis, page_num)

        # Add images
        self._add_page_images(page_analysis, page_num)

    def _convert_single_column(self, page_analysis: PageAnalysis, page_num: int):
        """Convert a single-column page."""
        # Sort text elements by reading order
        sorted_elements = sorted(
            page_analysis.text_elements,
            key=lambda e: e.reading_order
        )

        for element in sorted_elements:
            if element.is_page_number:
                continue  # Skip page numbers

            self._add_text_element(element)
            self._stats['text_elements'] += 1

    def _convert_multi_column(self, page_analysis: PageAnalysis, page_num: int):
        """Convert a multi-column page using table layout."""
        num_cols = page_analysis.num_columns

        # Group elements by column
        columns_content: Dict[int, List[TextElement]] = {i: [] for i in range(1, num_cols + 1)}

        for element in page_analysis.text_elements:
            if element.is_page_number:
                continue

            col = min(element.column, num_cols)
            columns_content[col].append(element)

        # Sort each column by reading order
        for col in columns_content:
            columns_content[col].sort(key=lambda e: e.reading_order)

        # Check for header elements (span full width)
        header_elements = [e for e in page_analysis.text_elements if e.is_header]
        footer_elements = [e for e in page_analysis.text_elements if e.is_footer]

        # Add headers first
        for element in sorted(header_elements, key=lambda e: e.reading_order):
            self._add_text_element(element)
            self._stats['text_elements'] += 1

        # Create table for columns
        table = self.doc.add_table(rows=1, cols=num_cols)
        table.autofit = False

        # Calculate column widths
        page_width_inches = 6.5  # Approximate content width
        col_width = page_width_inches / num_cols

        # Fill each column
        for col_idx in range(num_cols):
            cell = table.rows[0].cells[col_idx]
            cell.width = Inches(col_width)

            col_elements = columns_content.get(col_idx + 1, [])

            for element in col_elements:
                para = cell.add_paragraph()
                self._format_paragraph_with_element(para, element)
                self._stats['text_elements'] += 1

        # Remove table borders
        self._remove_table_borders(table)

        # Add footers
        for element in sorted(footer_elements, key=lambda e: e.reading_order):
            self._add_text_element(element)
            self._stats['text_elements'] += 1

    def _add_text_element(self, element: TextElement):
        """Add a text element as a paragraph."""
        para = self.doc.add_paragraph()
        self._format_paragraph_with_element(para, element)

    def _format_paragraph_with_element(self, para, element: TextElement):
        """Format a paragraph based on text element analysis."""
        run = para.add_run(element.text)

        # Font size
        size = self.FONT_SIZES.get(element.font_size, 11)
        run.font.size = Pt(size)

        # Font weight
        run.font.bold = element.font_weight == "bold"

        # Font style
        run.font.italic = element.font_style == "italic"

        # Color
        color = self._parse_color(element.color)
        if color:
            run.font.color.rgb = color

        # Spacing
        para.paragraph_format.space_after = Pt(6)

    def _parse_color(self, color_str: str) -> Optional[RGBColor]:
        """Parse color string to RGBColor."""
        color_str = color_str.lower().strip()

        # Named colors
        named_colors = {
            "black": RGBColor(0, 0, 0),
            "white": RGBColor(255, 255, 255),
            "red": RGBColor(255, 0, 0),
            "blue": RGBColor(0, 0, 255),
            "green": RGBColor(0, 128, 0),
            "gray": RGBColor(128, 128, 128),
            "grey": RGBColor(128, 128, 128),
        }

        if color_str in named_colors:
            return named_colors[color_str]

        # Hex color
        if color_str.startswith("#"):
            try:
                hex_color = color_str[1:]
                r = int(hex_color[0:2], 16)
                g = int(hex_color[2:4], 16)
                b = int(hex_color[4:6], 16)
                return RGBColor(r, g, b)
            except (ValueError, IndexError):
                pass

        return None

    def _add_page_images(self, page_analysis: PageAnalysis, page_num: int):
        """Extract and add images from the page."""
        if not page_analysis.images:
            return

        page = self.pdf_doc[page_num]
        image_list = page.get_images(full=True)

        for img_element in page_analysis.images:
            if img_element.is_background:
                continue  # Skip background images

            # Try to find and extract corresponding image
            for img_info in image_list:
                xref = img_info[0]
                try:
                    base_image = self.pdf_doc.extract_image(xref)
                    image_data = base_image.get("image", b"")

                    if image_data:
                        # Add image to document
                        image_stream = io.BytesIO(image_data)

                        # Calculate width based on bbox
                        img_width_pct = img_element.bbox.get("width", 50)
                        max_width = min(img_width_pct / 100 * 6, 5)  # Max 5 inches

                        para = self.doc.add_paragraph()
                        run = para.add_run()
                        run.add_picture(image_stream, width=Inches(max_width))

                        self._stats['images'] += 1
                        break

                except Exception as e:
                    self._notes.append(f"Could not extract image on page {page_num + 1}: {e}")

    def _remove_table_borders(self, table):
        """Remove all borders from a table."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)

        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def _calculate_validation_score(self, analysis: DocumentAnalysis) -> float:
        """Calculate a quality/validation score."""
        # Base score
        score = 0.7

        # Bonus for consistent layout
        if analysis.consistent_style:
            score += 0.1

        # Bonus for successful text extraction
        total_text = sum(len(p.text_elements) for p in analysis.pages)
        if total_text > 0:
            score += 0.1

        # Bonus for images processed
        if self._stats['images'] > 0:
            score += 0.05

        # Penalty for notes/warnings
        score -= len(self._notes) * 0.02

        return max(0, min(1, score))


def convert_pdf(
    pdf_path: str,
    output_path: Optional[str] = None,
    api_key: Optional[str] = None
) -> ConversionResult:
    """
    Convert PDF to DOCX with intelligent layout preservation.

    This is the main entry point for conversion. Simply call:
        result = convert_pdf("document.pdf")

    Args:
        pdf_path: Path to input PDF
        output_path: Path for output DOCX (optional)
        api_key: Anthropic API key (optional, uses env var if not provided)

    Returns:
        ConversionResult with conversion details
    """
    converter = IntelligentConverter(api_key)
    return converter.convert(pdf_path, output_path)

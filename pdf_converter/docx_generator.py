"""
DOCX Generator

Creates accurate DOCX from extracted PDF content.
Key features:
- Elements placed in exact reading order
- Proper table rendering with borders
- Images at correct positions
- Multi-column support using Word sections or tables
- Formatting preservation (fonts, sizes, colors)
- Accurate spacing and layout preservation
"""

import io
import re
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .pdf_extractor import (
    PageContent, TextBlock, ImageInfo, TableInfo, TableCell,
    PageElement, ElementType, ColumnInfo
)


@dataclass
class GenerationResult:
    """Result of DOCX generation."""
    success: bool
    output_path: str
    pages_processed: int
    text_blocks_written: int
    images_added: int
    tables_added: int
    errors: List[str]


class DOCXGenerator:
    """
    Generates DOCX from extracted PDF content.

    Creates accurate layout by:
    1. Processing elements in exact reading order
    2. Handling multi-column layouts with tables
    3. Placing images inline at correct positions
    4. Rendering tables with proper structure
    5. Preserving accurate spacing and line heights
    """

    # Font size limits
    MIN_FONT_SIZE = 6
    MAX_FONT_SIZE = 48

    # Spacing constants (in points)
    DEFAULT_LINE_SPACING_MULTIPLIER = 1.15  # Slightly more than single
    MIN_PARAGRAPH_SPACING = 0
    DEFAULT_PARAGRAPH_SPACING = 6
    MAX_PARAGRAPH_SPACING = 24

    # Layout constants
    CONTENT_WIDTH_INCHES = 7.0  # Total content width
    COLUMN_GAP_INCHES = 0.25  # Gap between columns
    
    def __init__(self):
        self.doc: Optional[Document] = None
        self._text_count = 0
        self._image_count = 0
        self._table_count = 0
        self._errors: List[str] = []
    
    def generate(
        self,
        pages: List[PageContent],
        layout_info=None,  # Optional, for backward compatibility
        output_path: str = "output.docx"
    ) -> GenerationResult:
        """
        Generate DOCX from extracted content.
        """
        self.doc = Document()
        self._text_count = 0
        self._image_count = 0
        self._table_count = 0
        self._errors = []
        
        if not pages:
            return GenerationResult(
                success=False,
                output_path=output_path,
                pages_processed=0,
                text_blocks_written=0,
                images_added=0,
                tables_added=0,
                errors=["No pages to process"]
            )
        
        # Setup document based on first page
        self._setup_document(pages[0])
        
        # Process each page
        for i, page in enumerate(pages):
            if i > 0:
                self._add_page_break()
            
            self._generate_page(page)
        
        # Save document
        try:
            self.doc.save(output_path)
            return GenerationResult(
                success=True,
                output_path=output_path,
                pages_processed=len(pages),
                text_blocks_written=self._text_count,
                images_added=self._image_count,
                tables_added=self._table_count,
                errors=self._errors
            )
        except Exception as e:
            return GenerationResult(
                success=False,
                output_path=output_path,
                pages_processed=len(pages),
                text_blocks_written=self._text_count,
                images_added=self._image_count,
                tables_added=self._table_count,
                errors=[f"Failed to save: {e}"]
            )
    
    def _setup_document(self, first_page: PageContent):
        """Setup document properties based on PDF page."""
        section = self.doc.sections[0]
        
        # Set page size (convert PDF points to DOCX units)
        section.page_width = Twips(first_page.width * 20)
        section.page_height = Twips(first_page.height * 20)
        
        # Set margins (narrower for better fidelity)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    
    def _add_page_break(self):
        """Add a page break."""
        self.doc.add_page_break()
    
    def _generate_page(self, page: PageContent):
        """Generate content for a single page."""
        if not page.elements:
            # Fallback: use text_blocks, images, tables directly
            self._generate_page_fallback(page)
            return
        
        # Check if multi-column
        if page.column_info.num_columns > 1:
            self._generate_multicolumn_page(page)
        else:
            self._generate_singlecolumn_page(page)
    
    def _generate_singlecolumn_page(self, page: PageContent):
        """Generate single-column page content."""
        prev_y = 0
        prev_elem_type = None
        prev_font_size = 11.0

        for elem in page.elements:
            # Calculate spacing based on vertical gap and element types
            spacing_pts = 0
            if prev_y > 0:
                gap = elem.y_position - prev_y
                # Get current element's font size for context
                curr_font_size = 11.0
                if elem.element_type == ElementType.TEXT:
                    curr_font_size = elem.element.primary_font_size

                # Calculate appropriate spacing based on gap size
                if gap > 0:
                    # Use font size as reference for spacing decisions
                    avg_font_size = (prev_font_size + curr_font_size) / 2

                    if gap < avg_font_size * 0.5:
                        # Very small gap - no extra spacing (same paragraph continuation)
                        spacing_pts = 0
                    elif gap < avg_font_size * 1.5:
                        # Normal line spacing within paragraph
                        spacing_pts = min(gap * 0.3, 6)
                    elif gap < avg_font_size * 3:
                        # Paragraph break
                        spacing_pts = min(gap * 0.5, 12)
                    else:
                        # Large gap - section break
                        spacing_pts = min(gap * 0.4, self.MAX_PARAGRAPH_SPACING)

                    # Add extra spacing for transitions between element types
                    if prev_elem_type != elem.element_type:
                        spacing_pts = max(spacing_pts, 6)

                if spacing_pts > 2:
                    self._add_vertical_space(spacing_pts)

            if elem.element_type == ElementType.TEXT:
                if not self._skip_header_footer_block(elem.element):
                    self._add_text_block(elem.element, page.width)
                prev_font_size = elem.element.primary_font_size
            elif elem.element_type == ElementType.IMAGE:
                self._add_image(elem.element, page.width)
                prev_font_size = 11.0
            elif elem.element_type == ElementType.TABLE:
                self._add_table(elem.element, page.width)
                prev_font_size = 10.0

            prev_y = elem.bbox[3]  # Bottom of element
            prev_elem_type = elem.element_type
    
    def _generate_multicolumn_page(self, page: PageContent):
        """
        Generate multi-column page content.
        
        Strategy: Group consecutive same-column elements into chunks,
        then render each chunk in a table-based layout.
        """
        # Separate elements by column and track order
        chunks = []
        current_chunk = {'columns': {}, 'type': 'columns', 'y_start': 0}
        
        for elem in page.elements:
            col = elem.column
            
            if col == 0:  # Full-width element
                # Save current chunk if not empty
                if any(current_chunk['columns'].values()):
                    chunks.append(current_chunk)
                
                # Add full-width element as its own chunk
                chunks.append({
                    'type': 'full_width',
                    'element': elem
                })
                
                # Start new column chunk
                current_chunk = {'columns': {}, 'type': 'columns', 'y_start': elem.bbox[3]}
            
            else:
                # Add to current column chunk
                if col not in current_chunk['columns']:
                    current_chunk['columns'][col] = []
                current_chunk['columns'][col].append(elem)
        
        # Don't forget the last chunk
        if any(current_chunk['columns'].values()):
            chunks.append(current_chunk)
        
        # Render each chunk
        for chunk in chunks:
            if chunk['type'] == 'full_width':
                elem = chunk['element']
                if elem.element_type == ElementType.TEXT:
                    if not self._skip_header_footer_block(elem.element):
                        self._add_text_block(elem.element, page.width)
                elif elem.element_type == ElementType.IMAGE:
                    self._add_image(elem.element, page.width)
                elif elem.element_type == ElementType.TABLE:
                    self._add_table(elem.element, page.width)
            
            elif chunk['type'] == 'columns':
                self._render_column_chunk(chunk['columns'], page)
    
    def _render_column_chunk(self, columns: Dict[int, List[PageElement]], page: PageContent):
        """Render a chunk of multi-column content using a table."""
        if not columns:
            return

        num_cols = page.column_info.num_columns

        # Create a table with the appropriate number of columns
        table = self.doc.add_table(rows=1, cols=num_cols)
        table.autofit = False

        # Calculate column widths based on actual column boundaries
        total_content_width = self.CONTENT_WIDTH_INCHES
        col_widths = []
        total_boundary_width = sum(x_end - x_start for x_start, x_end in page.column_info.boundaries)

        for x_start, x_end in page.column_info.boundaries:
            if total_boundary_width > 0:
                width_pct = (x_end - x_start) / total_boundary_width
            else:
                width_pct = 1.0 / num_cols
            # Subtract small amount for column gap
            col_width = width_pct * total_content_width - (self.COLUMN_GAP_INCHES / num_cols)
            col_widths.append(max(col_width, 1.0))  # Minimum 1 inch

        # Fill in the columns
        for col_idx in range(num_cols):
            cell = table.rows[0].cells[col_idx]

            if col_idx < len(col_widths):
                cell.width = Inches(col_widths[col_idx])

            # Set cell properties for clean layout with proper padding
            self._setup_cell(cell, add_padding=True)

            # Add content to this column
            col_num = col_idx + 1
            if col_num in columns:
                prev_y = 0
                for elem in columns[col_num]:
                    # Add spacing between elements
                    if prev_y > 0:
                        gap = elem.y_position - prev_y
                        if gap > 10:
                            # Add spacing paragraph
                            spacing_pts = min(gap * 0.3, 12)
                            if spacing_pts > 2:
                                p = cell.add_paragraph()
                                p.paragraph_format.space_after = Pt(spacing_pts)
                                p.paragraph_format.space_before = Pt(0)

                    if elem.element_type == ElementType.TEXT:
                        if not self._skip_header_footer_block(elem.element):
                            self._add_text_block_to_cell(cell, elem.element, page.width)
                    elif elem.element_type == ElementType.IMAGE:
                        self._add_image_to_cell(cell, elem.element, col_widths[col_idx] * 72)  # Convert to points

                    prev_y = elem.bbox[3]

        # Remove table borders for seamless column layout
        self._remove_table_borders(table)

        # Set table width
        self._set_table_width(table, total_content_width)
    
    def _setup_cell(self, cell, add_padding: bool = False):
        """Setup cell properties for clean layout."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Remove existing margins if present
        existing_mar = tcPr.find(qn('w:tcMar'))
        if existing_mar is not None:
            tcPr.remove(existing_mar)

        # Set cell margins
        tcMar = OxmlElement('w:tcMar')
        if add_padding:
            # Add small padding for better readability in multi-column
            margins = {
                'top': '57',      # ~1pt
                'left': '115',    # ~2pt
                'bottom': '57',   # ~1pt
                'right': '172'    # ~3pt (gutter side)
            }
        else:
            margins = {
                'top': '0',
                'left': '0',
                'bottom': '0',
                'right': '0'
            }

        for margin_name, value in margins.items():
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), value)
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)

        # Vertical alignment - top
        existing_valign = tcPr.find(qn('w:vAlign'))
        if existing_valign is not None:
            tcPr.remove(existing_valign)
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'top')
        tcPr.append(vAlign)

    def _set_table_width(self, table, width_inches: float):
        """Set the total width of a table."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

        # Remove existing width
        existing_width = tblPr.find(qn('w:tblW'))
        if existing_width is not None:
            tblPr.remove(existing_width)

        # Set new width
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(int(width_inches * 1440)))  # Convert to twips
        tblW.set(qn('w:type'), 'dxa')
        tblPr.insert(0, tblW)

        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
    
    def _generate_page_fallback(self, page: PageContent):
        """Fallback page generation using text_blocks, images, tables directly."""
        # Combine all elements with Y positions
        all_items = []
        
        for block in page.text_blocks:
            all_items.append(('text', block.y_position, block))
        
        for img in page.images:
            all_items.append(('image', img.y_position, img))
        
        for table in page.tables:
            all_items.append(('table', table.y_position, table))
        
        # Sort by Y position
        all_items.sort(key=lambda x: x[1])
        
        # Add each item
        prev_y = 0
        for item_type, y_pos, item in all_items:
            # Add spacing
            if prev_y > 0 and y_pos - prev_y > 15:
                self._add_vertical_space(min((y_pos - prev_y) / 3, 18))
            
            if item_type == 'text':
                if not self._skip_header_footer_block(item):
                    self._add_text_block(item, page.width)
            elif item_type == 'image':
                self._add_image(item, page.width)
            elif item_type == 'table':
                self._add_table(item, page.width)
            
            prev_y = y_pos
    
    def _skip_header_footer_block(self, block: TextBlock) -> bool:
        """Skip header/footer blocks that are only page numbers (e.g. '343', '344')."""
        if not block.is_header and not block.is_footer:
            return False
        text = block.text.strip()
        if not text or len(text) > 10:
            return False
        return text.isdigit()

    def _add_vertical_space(self, points: float):
        """Add vertical space between elements."""
        if points < 2:
            return
        # Clamp spacing to reasonable limits
        points = min(max(points, 2), self.MAX_PARAGRAPH_SPACING)
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(points)
        pf = para.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # Make the spacer paragraph minimal height
        run = para.add_run()
        run.font.size = Pt(1)
    
    def _add_text_block(self, block: TextBlock, page_width: float):
        """Add a text block as paragraph(s)."""
        # Create paragraph
        para = self.doc.add_paragraph()

        # Determine if this block should flow as continuous text or preserve line breaks
        should_flow = self._should_flow_text(block)

        # Add text with formatting
        for i, line in enumerate(block.lines):
            if i > 0:
                if should_flow:
                    # Check if previous line ends with hyphen (word continuation)
                    prev_text = block.lines[i - 1].text.rstrip()
                    if prev_text.endswith('-'):
                        # Don't add space, the hyphen indicates word continuation
                        pass
                    else:
                        # Add space between lines for flowing text
                        para.add_run(' ')
                else:
                    # Add soft line break for non-flowing text (preserves layout)
                    para.add_run('\n')

            # Handle line text - strip trailing whitespace but preserve leading
            line_text = line.text.rstrip()

            # If flowing text and line ends with hyphen, remove the hyphen
            if should_flow and line_text.endswith('-') and i < len(block.lines) - 1:
                line_text = line_text[:-1]

            run = para.add_run(line_text)

            # Apply formatting
            self._apply_run_formatting(run, line)

        # Paragraph formatting
        self._apply_paragraph_formatting(para, block, page_width)

        self._text_count += 1

    def _should_flow_text(self, block: TextBlock) -> bool:
        """Determine if text in block should flow (reflow) or preserve line breaks."""
        if not block.lines or len(block.lines) < 2:
            return False

        # Check if this looks like a paragraph (continuous prose) vs formatted text
        # Indicators of flowing text:
        # 1. Lines are roughly similar width (not ragged like a list)
        # 2. No special formatting patterns (bullets, numbers)
        # 3. Lines end with words (not punctuation that indicates line end)

        line_widths = [line.bbox[2] - line.bbox[0] for line in block.lines]
        if not line_widths:
            return False

        max_width = max(line_widths)
        if max_width == 0:
            return False

        # Calculate how many lines are "full width" (more than 85% of max)
        full_width_lines = sum(1 for w in line_widths[:-1] if w > max_width * 0.85)

        # If most lines are full width, it's likely flowing prose
        if len(block.lines) > 2 and full_width_lines >= (len(block.lines) - 1) * 0.6:
            # Additional check: first character patterns
            first_line_text = block.lines[0].text.strip()
            # Don't flow if it starts with bullet-like patterns
            if first_line_text and first_line_text[0] in '•●○■□▪▫-–—*+>':
                return False
            # Don't flow if it starts with numbers followed by punctuation (list)
            if re.match(r'^\d+[\.\)]\s', first_line_text):
                return False
            return True

        return False
    
    def _add_text_block_to_cell(self, cell, block: TextBlock, page_width: float = 612):
        """Add a text block to a table cell."""
        para = cell.add_paragraph()

        # Determine if this block should flow as continuous text
        should_flow = self._should_flow_text(block)

        for i, line in enumerate(block.lines):
            if i > 0:
                if should_flow:
                    prev_text = block.lines[i - 1].text.rstrip()
                    if not prev_text.endswith('-'):
                        para.add_run(' ')
                else:
                    para.add_run('\n')

            line_text = line.text.rstrip()
            if should_flow and line_text.endswith('-') and i < len(block.lines) - 1:
                line_text = line_text[:-1]

            run = para.add_run(line_text)
            self._apply_run_formatting(run, line)

        self._apply_paragraph_formatting(para, block, page_width)
        self._text_count += 1
    
    def _apply_run_formatting(self, run, line):
        """Apply formatting to a run based on TextLine properties."""
        # Font size - round to nearest 0.5 for cleaner output
        size = max(self.MIN_FONT_SIZE, min(self.MAX_FONT_SIZE, line.font_size))
        size = round(size * 2) / 2  # Round to nearest 0.5
        run.font.size = Pt(size)

        # Bold/Italic
        run.font.bold = line.is_bold
        run.font.italic = line.is_italic

        # Color - apply all colors including black for consistency
        if line.color:
            r, g, b = line.color
            # Only skip if it's exactly black (0, 0, 0) and we want Word default
            if r != 0 or g != 0 or b != 0:
                run.font.color.rgb = RGBColor(r, g, b)

        # Font name - always set for consistency
        font_name = self._clean_font_name(line.font_name)
        run.font.name = font_name

        # Set East Asian and Complex Script fonts too for full Unicode support
        rPr = run._r.get_or_add_rPr()

        # Set fonts for different scripts
        fonts = OxmlElement('w:rFonts')
        fonts.set(qn('w:ascii'), font_name)
        fonts.set(qn('w:hAnsi'), font_name)
        fonts.set(qn('w:eastAsia'), font_name)
        fonts.set(qn('w:cs'), font_name)

        # Remove existing rFonts if present
        existing_fonts = rPr.find(qn('w:rFonts'))
        if existing_fonts is not None:
            rPr.remove(existing_fonts)
        rPr.insert(0, fonts)
    
    def _apply_paragraph_formatting(self, para, block: TextBlock, page_width: float = 612):
        """Apply paragraph-level formatting."""
        pf = para.paragraph_format

        # Calculate spacing based on font size
        font_size = block.primary_font_size
        space_after = min(font_size * 0.3, 6)  # Proportional spacing

        pf.space_before = Pt(0)
        pf.space_after = Pt(space_after)

        # Line spacing - use multiple for better readability
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = self.DEFAULT_LINE_SPACING_MULTIPLIER

        # Detect alignment from block position
        block_center = (block.bbox[0] + block.bbox[2]) / 2
        page_center = page_width / 2

        # Calculate margins (approximate content area)
        margin = page_width * 0.08  # ~8% margins on each side
        content_left = margin
        content_right = page_width - margin
        content_center = page_width / 2

        # Determine alignment based on position
        block_left = block.bbox[0]
        block_right = block.bbox[2]
        block_width = block_right - block_left

        # Check if block is centered
        left_margin_dist = abs(block_left - content_left)
        right_margin_dist = abs(block_right - content_right)

        if block_width < (content_right - content_left) * 0.6:
            # Short block - check if centered
            if abs(left_margin_dist - right_margin_dist) < 20:
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif right_margin_dist < left_margin_dist * 0.5:
                # Block is closer to right margin
                pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Wide block - justify if it spans most of the width
            if block_width > (content_right - content_left) * 0.9:
                pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _clean_font_name(self, font_name: str) -> str:
        """Clean PDF font name for Word with comprehensive font mapping."""
        if not font_name:
            return "Calibri"  # Default to Calibri instead of empty

        # Remove subset prefix (e.g., "ABCDEF+")
        if "+" in font_name:
            font_name = font_name.split("+", 1)[-1]

        # Normalize: remove hyphens and make lowercase for matching
        font_lower = font_name.lower().replace("-", "").replace(" ", "")

        # Comprehensive font family mappings
        font_families = {
            # Arial family
            "arial": "Arial",
            "arialmt": "Arial",
            "arialbold": "Arial",
            "arialitalic": "Arial",
            "arialbolditalic": "Arial",
            "arialboldmt": "Arial",
            "arialitalicmt": "Arial",
            "arialbolditalicmt": "Arial",
            "arialnarrow": "Arial Narrow",
            "arialblack": "Arial Black",

            # Helvetica -> Arial (common substitution)
            "helvetica": "Arial",
            "helveticabold": "Arial",
            "helveticaoblique": "Arial",
            "helveticaboldoblique": "Arial",
            "helveticaneue": "Arial",
            "helveticaneuebold": "Arial",
            "helveticaneuelight": "Arial",

            # Times family
            "times": "Times New Roman",
            "timesnewroman": "Times New Roman",
            "timesnewromanpsmt": "Times New Roman",
            "timesnewromanps": "Times New Roman",
            "timesnewromanbold": "Times New Roman",
            "timesnewromanitalic": "Times New Roman",
            "timesnewromanbolditalic": "Times New Roman",
            "timesroman": "Times New Roman",

            # Courier family
            "courier": "Courier New",
            "couriernew": "Courier New",
            "couriernewpsmt": "Courier New",
            "couriernewtps": "Courier New",

            # Calibri family (Office default)
            "calibri": "Calibri",
            "calibribold": "Calibri",
            "calibrilight": "Calibri Light",

            # Cambria family
            "cambria": "Cambria",
            "cambriabold": "Cambria",
            "cambriamath": "Cambria Math",

            # Georgia family
            "georgia": "Georgia",
            "georgiabold": "Georgia",
            "georgiaitalic": "Georgia",

            # Verdana family
            "verdana": "Verdana",
            "verdanabold": "Verdana",
            "verdanaitalic": "Verdana",

            # Tahoma family
            "tahoma": "Tahoma",
            "tahomabold": "Tahoma",

            # Trebuchet family
            "trebuchetms": "Trebuchet MS",
            "trebuchet": "Trebuchet MS",

            # Comic Sans
            "comicsansms": "Comic Sans MS",
            "comicsans": "Comic Sans MS",

            # Impact
            "impact": "Impact",

            # Garamond family
            "garamond": "Garamond",
            "garamondbold": "Garamond",

            # Palatino family
            "palatino": "Palatino Linotype",
            "palatinolinotype": "Palatino Linotype",
            "bookantiqua": "Book Antiqua",

            # Century family
            "centurygothic": "Century Gothic",
            "centuryschoolbook": "Century Schoolbook",

            # Segoe family (Windows)
            "segoeui": "Segoe UI",
            "segoeuibold": "Segoe UI",
            "segoeuisemibold": "Segoe UI Semibold",
            "segoeuilight": "Segoe UI Light",

            # Consolas (monospace)
            "consolas": "Consolas",

            # Symbol fonts
            "symbol": "Symbol",
            "wingdings": "Wingdings",
            "webdings": "Webdings",
            "zapfdingbats": "Wingdings",

            # Open Sans
            "opensans": "Arial",
            "opensansbold": "Arial",
            "opensanslight": "Arial",

            # Roboto -> Arial
            "roboto": "Arial",
            "robotobold": "Arial",
            "robotolight": "Arial",

            # Source Sans -> Arial
            "sourcesanspro": "Arial",
            "sourcesans": "Arial",

            # Lato -> Arial
            "lato": "Arial",
            "latobold": "Arial",

            # Montserrat -> Arial
            "montserrat": "Arial",
            "montserratbold": "Arial",
        }

        # Try exact match first
        if font_lower in font_families:
            return font_families[font_lower]

        # Try to match font family by checking if the font name contains known families
        font_family_prefixes = [
            ("arial", "Arial"),
            ("helvetica", "Arial"),
            ("times", "Times New Roman"),
            ("courier", "Courier New"),
            ("calibri", "Calibri"),
            ("cambria", "Cambria"),
            ("georgia", "Georgia"),
            ("verdana", "Verdana"),
            ("tahoma", "Tahoma"),
            ("trebuchet", "Trebuchet MS"),
            ("garamond", "Garamond"),
            ("palatino", "Palatino Linotype"),
            ("segoe", "Segoe UI"),
            ("consolas", "Consolas"),
            ("roboto", "Arial"),
            ("opensans", "Arial"),
            ("lato", "Arial"),
            ("montserrat", "Arial"),
        ]

        for prefix, mapped_font in font_family_prefixes:
            if prefix in font_lower:
                return mapped_font

        # Remove common suffixes and try again
        suffixes_to_remove = [
            "bold", "italic", "oblique", "light", "medium", "semibold",
            "regular", "book", "condensed", "extended", "narrow",
            "mt", "ps", "psmt", "std", "pro"
        ]

        cleaned = font_lower
        for suffix in suffixes_to_remove:
            if cleaned.endswith(suffix):
                cleaned = cleaned[:-len(suffix)]

        if cleaned in font_families:
            return font_families[cleaned]

        for prefix, mapped_font in font_family_prefixes:
            if prefix in cleaned:
                return mapped_font

        # If nothing matches, return Calibri as safe default
        return "Calibri"
    
    def _add_image(self, img: ImageInfo, page_width: float):
        """Add an image to the document."""
        if not img.data:
            return

        try:
            # Try to convert image data to a format python-docx can handle
            image_data = self._prepare_image_data(img.data, img.ext)
            if not image_data:
                self._errors.append("Image: Could not prepare image data")
                return

            image_stream = io.BytesIO(image_data)

            # Calculate size - preserve aspect ratio, fit within page
            img_width_ratio = img.width / page_width if page_width > 0 else 0.5

            # Scale to fit content width while preserving proportions
            max_width_inches = min(img_width_ratio * self.CONTENT_WIDTH_INCHES, self.CONTENT_WIDTH_INCHES)
            max_width_inches = max(max_width_inches, 0.5)  # Min 0.5 inch

            # Ensure aspect ratio is preserved
            if img.width > 0 and img.height > 0:
                aspect_ratio = img.height / img.width
                max_height_inches = max_width_inches * aspect_ratio

                # Limit height to reasonable value (8 inches)
                if max_height_inches > 8.0:
                    max_height_inches = 8.0
                    max_width_inches = max_height_inches / aspect_ratio

            # Create paragraph and add image
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = para.add_run()
            run.add_picture(image_stream, width=Inches(max_width_inches))

            # Proportional spacing based on image size
            spacing = min(max(max_width_inches * 2, 4), 12)
            para.paragraph_format.space_before = Pt(spacing)
            para.paragraph_format.space_after = Pt(spacing)

            self._image_count += 1

        except Exception as e:
            self._errors.append(f"Image error: {e}")

    def _prepare_image_data(self, data: bytes, ext: str) -> Optional[bytes]:
        """Prepare image data for insertion into DOCX, converting if necessary."""
        if not data:
            return None

        try:
            # Try to use PIL/Pillow to validate and potentially convert the image
            from PIL import Image

            img_stream = io.BytesIO(data)
            try:
                pil_image = Image.open(img_stream)
                pil_image.load()  # Force load to verify it's valid

                # Convert to RGB if necessary (handles CMYK, etc.)
                if pil_image.mode in ('CMYK', 'P', 'LA', 'PA'):
                    pil_image = pil_image.convert('RGB')
                elif pil_image.mode == 'RGBA':
                    # Keep RGBA for transparency support
                    pass
                elif pil_image.mode != 'RGB':
                    pil_image = pil_image.convert('RGB')

                # Save to PNG format which is well supported
                output = io.BytesIO()
                save_format = 'PNG' if pil_image.mode == 'RGBA' else 'JPEG'
                pil_image.save(output, format=save_format, quality=95)
                return output.getvalue()

            except Exception:
                # PIL couldn't process it, return original data
                pass

        except ImportError:
            # PIL not available, try using original data
            pass

        # Return original data as fallback
        return data
    
    def _add_image_to_cell(self, cell, img: ImageInfo, max_width_pts: float):
        """Add an image to a table cell."""
        if not img.data:
            return

        try:
            # Prepare image data
            image_data = self._prepare_image_data(img.data, img.ext)
            if not image_data:
                return

            image_stream = io.BytesIO(image_data)

            # Convert max_width from points to inches
            max_width_inches = max_width_pts / 72 if max_width_pts > 0 else 3.0

            # Calculate size for cell while preserving aspect ratio
            if img.width > 0:
                img_width_inches = img.width / 72  # Convert points to inches
                width_inches = min(img_width_inches, max_width_inches * 0.9)  # 90% of cell width max
            else:
                width_inches = max_width_inches * 0.5

            width_inches = max(width_inches, 0.5)  # Min 0.5 inch

            para = cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(image_stream, width=Inches(width_inches))

            # Minimal spacing
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)

            self._image_count += 1

        except Exception as e:
            self._errors.append(f"Cell image error: {e}")
    
    def _add_table(self, table_info: TableInfo, page_width: float):
        """Add a proper table to the document."""
        if not table_info.cells or table_info.num_rows == 0:
            return
        
        try:
            # Create table
            table = self.doc.add_table(rows=table_info.num_rows, cols=table_info.num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Calculate width
            table_width_ratio = (table_info.bbox[2] - table_info.bbox[0]) / page_width
            table_width_inches = min(table_width_ratio * 7.5, 7.0)
            col_width = table_width_inches / table_info.num_cols
            
            # Fill cells
            for row_idx, row_cells in enumerate(table_info.cells):
                for col_idx, cell_info in enumerate(row_cells):
                    if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.width = Inches(col_width)
                        
                        # Set cell text
                        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        
                        # Clear any default text
                        if para.runs:
                            para.clear()
                        
                        run = para.add_run(cell_info.text)
                        
                        # Formatting
                        run.font.size = Pt(cell_info.font_size)
                        run.font.bold = cell_info.is_bold
                        
                        # Cell vertical alignment
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            
            # Apply table borders
            self._style_table_borders(table)
            
            # Spacing around table
            # Add empty paragraph before if needed
            
            self._table_count += 1
            
        except Exception as e:
            self._errors.append(f"Table error: {e}")
    
    def _style_table_borders(self, table):
        """Add borders to table."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
    
    def _remove_table_borders(self, table):
        """Remove borders from table (for column layout)."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

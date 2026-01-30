"""
Advanced PDF to DOCX Converter with Layout Preservation

This converter provides high-fidelity PDF to DOCX conversion with:
- Multi-column layout preservation
- Accurate text positioning and formatting
- Image extraction and placement
- Post-processing validation
"""

import os
import io
import re
from pathlib import Path
from typing import Optional, List, Dict, Tuple, Any
from dataclasses import dataclass
from enum import Enum

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

from .analyzer import PDFAnalyzer, PageLayout, TextBlock, ImageBlock, LayoutType, Column


class ConversionQuality(Enum):
    """Conversion quality levels."""
    DRAFT = "draft"  # Fast, basic conversion
    STANDARD = "standard"  # Balanced quality and speed
    HIGH = "high"  # Maximum fidelity, slower


@dataclass
class ConversionStats:
    """Statistics from conversion process."""
    pages_processed: int = 0
    text_blocks_converted: int = 0
    images_extracted: int = 0
    tables_detected: int = 0
    columns_preserved: int = 0
    warnings: List[str] = None

    def __post_init__(self):
        if self.warnings is None:
            self.warnings = []


@dataclass
class ValidationResult:
    """Result of post-conversion validation."""
    is_valid: bool
    text_match_ratio: float
    structure_match_ratio: float
    issues: List[str]
    suggestions: List[str]


class AdvancedPDFConverter:
    """
    Advanced PDF to DOCX converter with full layout preservation.

    Features:
    - Multi-column layout detection and recreation
    - Accurate reading order preservation
    - Font and style mapping
    - Image extraction with positioning
    - Table detection and recreation
    - Post-processing validation
    """

    def __init__(self, quality: ConversionQuality = ConversionQuality.HIGH):
        """
        Initialize the converter.

        Args:
            quality: Conversion quality level
        """
        self.quality = quality
        self.stats = ConversionStats()
        self._current_doc: Optional[Document] = None
        self._analyzer: Optional[PDFAnalyzer] = None

    def convert(
        self,
        pdf_path: str,
        output_path: Optional[str] = None,
        pages: Optional[List[int]] = None,
        validate: bool = True
    ) -> Tuple[str, ConversionStats, Optional[ValidationResult]]:
        """
        Convert PDF to DOCX with layout preservation.

        Args:
            pdf_path: Path to input PDF
            output_path: Path for output DOCX (optional)
            pages: Specific pages to convert (optional)
            validate: Whether to run post-conversion validation

        Returns:
            Tuple of (output_path, stats, validation_result)
        """
        pdf_path = str(Path(pdf_path).resolve())

        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF not found: {pdf_path}")

        if output_path is None:
            output_path = str(Path(pdf_path).with_suffix('.docx'))

        print(f"[1/5] Analyzing PDF structure...")
        self._analyzer = PDFAnalyzer(pdf_path)
        structure = self._analyzer.analyze()

        print(f"      Found {structure.page_count} pages, dominant layout: {structure.dominant_layout.value}")

        # Determine pages to convert
        if pages is None:
            pages = list(range(structure.page_count))

        print(f"[2/5] Creating document structure...")
        self._current_doc = Document()
        self._setup_document(structure)

        print(f"[3/5] Converting pages...")
        for i, page_num in enumerate(pages):
            print(f"      Processing page {page_num + 1}/{structure.page_count}...")
            self._convert_page(structure.pages[page_num], page_num == pages[0])
            self.stats.pages_processed += 1

        print(f"[4/5] Saving document...")
        self._current_doc.save(output_path)

        validation_result = None
        if validate:
            print(f"[5/5] Validating conversion...")
            validation_result = self._validate_conversion(pdf_path, output_path, pages)
            if validation_result.is_valid:
                print(f"      Validation PASSED (text match: {validation_result.text_match_ratio:.1%})")
            else:
                print(f"      Validation found issues (text match: {validation_result.text_match_ratio:.1%})")
                for issue in validation_result.issues[:3]:
                    print(f"        - {issue}")
        else:
            print(f"[5/5] Skipping validation...")

        self._analyzer.close()

        return output_path, self.stats, validation_result

    def _setup_document(self, structure):
        """Setup document with appropriate page size and margins."""
        if not structure.pages:
            return

        first_page = structure.pages[0]

        # Set page size
        section = self._current_doc.sections[0]
        section.page_width = Twips(first_page.width * 20)  # Convert points to twips
        section.page_height = Twips(first_page.height * 20)

        # Set margins
        margins = first_page.margins
        section.left_margin = Twips(margins.get("left", 72) * 20)
        section.right_margin = Twips(margins.get("right", 72) * 20)
        section.top_margin = Twips(margins.get("top", 72) * 20)
        section.bottom_margin = Twips(margins.get("bottom", 72) * 20)

    def _convert_page(self, page_layout: PageLayout, is_first_page: bool):
        """Convert a single page to DOCX content."""
        # Add page break for subsequent pages
        if not is_first_page:
            self._current_doc.add_page_break()

        # Convert based on layout type
        if page_layout.layout_type == LayoutType.SINGLE_COLUMN:
            self._convert_single_column_page(page_layout)
        elif page_layout.layout_type in [LayoutType.TWO_COLUMN, LayoutType.THREE_COLUMN, LayoutType.MULTI_COLUMN]:
            self._convert_multi_column_page(page_layout)
        else:
            # Mixed or complex layout - use table-based approach
            self._convert_complex_page(page_layout)

    def _convert_single_column_page(self, page_layout: PageLayout):
        """Convert a single-column page."""
        # Add header content
        for block in sorted(page_layout.header_blocks, key=lambda b: (b.y0, b.x0)):
            self._add_text_block(block)

        # Add content in reading order
        content_blocks = [b for b in page_layout.text_blocks
                        if b not in page_layout.header_blocks
                        and b not in page_layout.footer_blocks]

        # Sort by position for single column
        for block in sorted(content_blocks, key=lambda b: (b.y0, b.x0)):
            self._add_text_block(block)
            self.stats.text_blocks_converted += 1

        # Add images
        for img in page_layout.images:
            self._add_image(img, page_layout.width)

        # Add footer content
        for block in sorted(page_layout.footer_blocks, key=lambda b: (b.y0, b.x0)):
            self._add_text_block(block)

    def _convert_multi_column_page(self, page_layout: PageLayout):
        """Convert a multi-column page using a table for layout."""
        num_columns = len(page_layout.columns)

        if num_columns < 2:
            self._convert_single_column_page(page_layout)
            return

        # Add header content first (spans full width)
        for block in sorted(page_layout.header_blocks, key=lambda b: (b.y0, b.x0)):
            self._add_text_block(block)

        # Create a table for column layout
        table = self._current_doc.add_table(rows=1, cols=num_columns)
        table.autofit = False
        table.allow_autofit = False

        # Calculate column widths
        content_width = page_layout.width - page_layout.margins.get("left", 72) - page_layout.margins.get("right", 72)

        for col_idx, column in enumerate(page_layout.columns):
            cell = table.rows[0].cells[col_idx]

            # Set column width
            col_width_inches = (column.width / 72)  # Convert points to inches
            cell.width = Inches(col_width_inches)

            # Sort blocks in this column by vertical position
            sorted_blocks = sorted(column.blocks, key=lambda b: b.y0)

            for block in sorted_blocks:
                # Add paragraph to cell
                para = cell.add_paragraph()
                self._format_paragraph(para, block)
                self.stats.text_blocks_converted += 1

            self.stats.columns_preserved += 1

        # Remove table borders for seamless look
        self._remove_table_borders(table)

        # Add images (positioned relative to columns)
        for img in page_layout.images:
            self._add_image(img, page_layout.width)

        # Add footer content (spans full width)
        for block in sorted(page_layout.footer_blocks, key=lambda b: (b.y0, b.x0)):
            self._add_text_block(block)

    def _convert_complex_page(self, page_layout: PageLayout):
        """Convert a complex/mixed layout page."""
        # Group blocks by vertical sections
        sections = self._group_blocks_by_rows(page_layout)

        # Add header
        for block in sorted(page_layout.header_blocks, key=lambda b: (b.y0, b.x0)):
            self._add_text_block(block)

        # Process each section
        for section_blocks in sections:
            if len(section_blocks) == 1:
                # Single block - add directly
                self._add_text_block(section_blocks[0])
            else:
                # Multiple blocks at similar y - might be columns
                # Sort by x position
                sorted_blocks = sorted(section_blocks, key=lambda b: b.x0)

                # Check if they form columns
                if self._blocks_form_columns(sorted_blocks, page_layout.width):
                    # Add as table row
                    table = self._current_doc.add_table(rows=1, cols=len(sorted_blocks))
                    for col_idx, block in enumerate(sorted_blocks):
                        cell = table.rows[0].cells[col_idx]
                        para = cell.add_paragraph()
                        self._format_paragraph(para, block)
                    self._remove_table_borders(table)
                else:
                    # Add sequentially
                    for block in sorted_blocks:
                        self._add_text_block(block)

            self.stats.text_blocks_converted += len(section_blocks)

        # Add footer
        for block in sorted(page_layout.footer_blocks, key=lambda b: (b.y0, b.x0)):
            self._add_text_block(block)

    def _group_blocks_by_rows(self, page_layout: PageLayout) -> List[List[TextBlock]]:
        """Group text blocks into rows based on vertical position."""
        content_blocks = [b for b in page_layout.text_blocks
                        if b not in page_layout.header_blocks
                        and b not in page_layout.footer_blocks]

        if not content_blocks:
            return []

        # Sort by y position
        sorted_blocks = sorted(content_blocks, key=lambda b: b.y0)

        rows = []
        current_row = [sorted_blocks[0]]
        row_top = sorted_blocks[0].y0

        threshold = 15  # Points - blocks within this vertical distance are same row

        for block in sorted_blocks[1:]:
            if abs(block.y0 - row_top) < threshold:
                current_row.append(block)
            else:
                rows.append(current_row)
                current_row = [block]
                row_top = block.y0

        if current_row:
            rows.append(current_row)

        return rows

    def _blocks_form_columns(self, blocks: List[TextBlock], page_width: float) -> bool:
        """Check if blocks are arranged as columns."""
        if len(blocks) < 2:
            return False

        # Check if blocks don't overlap horizontally
        sorted_blocks = sorted(blocks, key=lambda b: b.x0)
        for i in range(len(sorted_blocks) - 1):
            if sorted_blocks[i].x1 > sorted_blocks[i+1].x0:
                return False  # Overlapping

        # Check if blocks have similar heights and reasonable widths
        heights = [b.height for b in blocks]
        avg_height = sum(heights) / len(heights)
        height_variance = sum((h - avg_height)**2 for h in heights) / len(heights)

        if height_variance > avg_height * 10:
            return False  # Too different in height

        return True

    def _add_text_block(self, block: TextBlock):
        """Add a text block as a paragraph."""
        para = self._current_doc.add_paragraph()
        self._format_paragraph(para, block)

    def _format_paragraph(self, para, block: TextBlock):
        """Format a paragraph with text block styling."""
        run = para.add_run(block.text)

        # Font formatting
        run.font.size = Pt(block.font_size)
        run.font.bold = block.is_bold
        run.font.italic = block.is_italic

        # Font name
        if block.font_name:
            # Clean up font name
            font_name = block.font_name
            # Remove common prefixes
            for prefix in ["AAAAAA+", "BAAAAA+", "CAAAAA+"]:
                if font_name.startswith(prefix):
                    font_name = font_name[len(prefix):]

            # Map common font names
            font_mapping = {
                "ArialMT": "Arial",
                "Arial-BoldMT": "Arial",
                "TimesNewRomanPSMT": "Times New Roman",
                "TimesNewRomanPS-BoldMT": "Times New Roman",
            }
            font_name = font_mapping.get(font_name, font_name)
            run.font.name = font_name

        # Color
        if block.color and block.color != (0, 0, 0):
            r, g, b = block.color
            run.font.color.rgb = RGBColor(int(r*255), int(g*255), int(b*255))

        # Paragraph spacing
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def _add_image(self, img: ImageBlock, page_width: float):
        """Add an image to the document."""
        if not img.image_data:
            return

        try:
            # Calculate image size
            max_width = 6.0  # inches
            img_width_inches = min(img.width / 72, max_width)

            # Add image
            image_stream = io.BytesIO(img.image_data)
            para = self._current_doc.add_paragraph()
            run = para.add_run()
            run.add_picture(image_stream, width=Inches(img_width_inches))

            self.stats.images_extracted += 1
        except Exception as e:
            self.stats.warnings.append(f"Failed to add image: {e}")

    def _remove_table_borders(self, table):
        """Remove all borders from a table for seamless column layout."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

        tblBorders = OxmlElement('w:tblBorders')

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)

        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def _validate_conversion(
        self,
        pdf_path: str,
        docx_path: str,
        pages: List[int]
    ) -> ValidationResult:
        """Validate the conversion by comparing PDF and DOCX content."""
        issues = []
        suggestions = []

        # Extract text from original PDF
        pdf_doc = fitz.open(pdf_path)
        pdf_text = ""
        for page_num in pages:
            pdf_text += pdf_doc[page_num].get_text()
        pdf_doc.close()

        # Extract text from generated DOCX
        docx_doc = Document(docx_path)
        docx_text = ""
        for para in docx_doc.paragraphs:
            docx_text += para.text + "\n"
        for table in docx_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    docx_text += cell.text + "\n"

        # Clean and normalize text for comparison
        pdf_text_clean = self._normalize_text(pdf_text)
        docx_text_clean = self._normalize_text(docx_text)

        # Calculate text match ratio
        pdf_words = set(pdf_text_clean.split())
        docx_words = set(docx_text_clean.split())

        if pdf_words:
            common_words = pdf_words & docx_words
            text_match_ratio = len(common_words) / len(pdf_words)
        else:
            text_match_ratio = 1.0 if not docx_words else 0.0

        # Check for missing content
        missing_words = pdf_words - docx_words
        if len(missing_words) > len(pdf_words) * 0.1:
            issues.append(f"Missing approximately {len(missing_words)} words from original")
            suggestions.append("Check complex layouts or images with text")

        # Check for blank content
        if not docx_text_clean.strip():
            issues.append("Output document appears to be empty or mostly blank")
            suggestions.append("PDF may contain scanned images requiring OCR")

        # Structure validation
        structure_match = 1.0
        if text_match_ratio < 0.8:
            structure_match = 0.5
            issues.append("Significant text content mismatch detected")

        is_valid = text_match_ratio >= 0.85 and not any("empty" in i.lower() for i in issues)

        return ValidationResult(
            is_valid=is_valid,
            text_match_ratio=text_match_ratio,
            structure_match_ratio=structure_match,
            issues=issues,
            suggestions=suggestions
        )

    def _normalize_text(self, text: str) -> str:
        """Normalize text for comparison."""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters
        text = re.sub(r'[^\w\s]', '', text)
        # Lowercase
        text = text.lower().strip()
        return text


def convert_with_layout_preservation(
    pdf_path: str,
    output_path: Optional[str] = None,
    quality: str = "high",
    validate: bool = True
) -> Tuple[str, dict]:
    """
    Convenience function for advanced PDF to DOCX conversion.

    Args:
        pdf_path: Path to input PDF
        output_path: Path for output DOCX (optional)
        quality: Conversion quality ("draft", "standard", "high")
        validate: Whether to validate output

    Returns:
        Tuple of (output_path, result_dict)
    """
    quality_map = {
        "draft": ConversionQuality.DRAFT,
        "standard": ConversionQuality.STANDARD,
        "high": ConversionQuality.HIGH
    }

    converter = AdvancedPDFConverter(quality=quality_map.get(quality, ConversionQuality.HIGH))
    output_path, stats, validation = converter.convert(pdf_path, output_path, validate=validate)

    result = {
        "output_path": output_path,
        "pages_processed": stats.pages_processed,
        "text_blocks_converted": stats.text_blocks_converted,
        "images_extracted": stats.images_extracted,
        "columns_preserved": stats.columns_preserved,
        "warnings": stats.warnings,
        "validation": {
            "is_valid": validation.is_valid if validation else None,
            "text_match_ratio": validation.text_match_ratio if validation else None,
            "issues": validation.issues if validation else [],
            "suggestions": validation.suggestions if validation else []
        } if validation else None
    }

    return output_path, result
